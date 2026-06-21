from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from datetime import datetime
from typing import Optional

from config import EXPORTS_DIR
from models import TranscriptionResult


def format_timestamp(seconds: float) -> str:
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    if h > 0:
        return f"{h:02d}:{m:02d}:{s:02d}"
    return f"{m:02d}:{s:02d}"


def generate_docx(
    result: TranscriptionResult, 
    original_filename: str, 
    refined_text: Optional[str] = None,
    refined_mode: Optional[str] = None
) -> Path:
    """Generate DOCX document from transcription result.
    
    Args:
        result: TranscriptionResult from Whisper
        original_filename: Name of the original audio file
        refined_text: Optional refined text to include instead of original
        refined_mode: Optional name of refinement mode used
        
    Returns:
        Path to generated DOCX file
    """
    doc = Document()

    # Add title
    title_text = f"Transcription: {original_filename}"
    if refined_mode:
        title_text += f" ({refined_mode} refined)"
    title = doc.add_heading(title_text, level=1)
    
    # Add metadata
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Model: {result.model}")
    if refined_mode:
        doc.add_paragraph(f"Refined with: {refined_mode}")
    doc.add_paragraph("")

    # Add content based on whether we have refined text
    text_to_export = refined_text or result.text
    
    if refined_text:
        # For refined text, just show the full refined content
        p = doc.add_paragraph()
        text_run = p.add_run(text_to_export)
        text_run.font.size = Pt(11)
    else:
        # For original transcription, show segments with timestamps
        for seg in result.segments:
            p = doc.add_paragraph()
            ts_run = p.add_run(f"[{format_timestamp(seg.start)} - {format_timestamp(seg.end)}]  ")
            ts_run.bold = True
            ts_run.font.size = Pt(9)
            ts_run.font.color.rgb = RGBColor(100, 100, 100)
            text_run = p.add_run(seg.text.strip())
            text_run.font.size = Pt(11)

    export_name = f"{Path(original_filename).stem}_transcription.docx"
    if refined_mode:
        export_name = f"{Path(original_filename).stem}_{refined_mode}_refined.docx"
    
    export_path = EXPORTS_DIR / export_name
    doc.save(str(export_path))
    return export_path
